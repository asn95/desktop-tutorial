"""
Generate Form3_Design.docx for C3MR Capstone Design Project
Follows the F300 alumni document format exactly.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

BASE = "/workspace/desktop-tutorial"

# ── helpers ──────────────────────────────────────────────────────────────────

def set_font(run, name="Times New Roman", size=12, bold=False, italic=False,
             color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def para(doc, text="", bold=False, italic=False, center=False,
         size=12, space_before=0, space_after=6, font="Times New Roman",
         color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        r = p.add_run(text)
        set_font(r, font, size, bold, italic, color)
    return p

def heading(doc, text, level=1, center=True, size=None, space_before=12):
    sizes = {1: 14, 2: 13, 3: 12, 4: 12}
    sz = size or sizes.get(level, 12)
    return para(doc, text, bold=True, center=center, size=sz,
                space_before=space_before, space_after=6)

def h2(doc, text):
    return heading(doc, text, level=2, center=True, size=12)

def h3(doc, text):
    return heading(doc, text, level=3, center=False, size=12)

def body(doc, text, indent=False):
    p = para(doc, text, size=12, space_before=0, space_after=6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    return p

def page_break(doc):
    doc.add_page_break()

def add_image(doc, path, width=Inches(5.5), caption=None):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        if caption:
            cp = doc.add_paragraph(caption)
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = cp.runs[0]
            set_font(r, size=11, bold=True)
    else:
        para(doc, f"[Figure: {caption or os.path.basename(path)}]",
             italic=True, center=True)

def simple_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cell.paragraphs[0].runs[0]
        set_font(r, bold=True, size=11)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if cell.paragraphs[0].runs:
                set_font(cell.paragraphs[0].runs[0], size=11)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if col_widths:
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                if ci < len(col_widths):
                    cell.width = col_widths[ci]
    return t

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size=12)
    return p

def numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_font(r, size=12)
    return p

# ── document ──────────────────────────────────────────────────────────────────

doc = Document()

# Margins
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(3.0)
    sec.right_margin  = Cm(2.5)

# Default Normal style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)

# ═══════════════════════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════════════════════
para(doc, space_before=0, space_after=0)  # top spacing

# University logo
logo = os.path.join(BASE, "uml", "images", "01_system_architecture.png")
# Try to use a generic placeholder; skip if not found
# (User can insert logo manually)
para(doc, "[Insert University Logo Here]", italic=True, center=True, size=11, space_after=12)

para(doc, "F300 Document", bold=True, center=True, size=16, space_after=2)
para(doc, "Capstone Design", bold=True, center=True, size=16, space_after=2)
para(doc, "Project Design", bold=True, center=True, size=16, space_after=8)
para(doc,
     "Integrated Operational Management System for C3MR:\n"
     "Web Admin Portal & Telegram Mini App Ecosystem",
     bold=True, center=True, size=14, space_after=16)

# Group member label
para(doc, "GROUP MEMBER:", bold=True, center=True, size=12, space_after=4)

# Group member table
members_tbl = doc.add_table(rows=4, cols=4)
members_tbl.style = "Table Grid"
members_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

col_names = ["No.", "Student Name", "Student ID", "Role"]
for ci, cn in enumerate(col_names):
    c = members_tbl.rows[0].cells[ci]
    c.text = cn
    c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(c.paragraphs[0].runs[0], bold=True, size=11)

members_data = [
    ("1.", "Auza Syamil Nabawi",       "001202300150", "Web Admin Portal & Backend"),
    ("2.", "Rashad Abdul Faqih",        "001202300149", "Telegram Mini App & Bot"),
    ("3.", "Atthariqul Hazam Albanna",  "012202300122", "System Architecture & Integration"),
]
for ri, (no, name, sid, role) in enumerate(members_data):
    row = members_tbl.rows[ri + 1]
    for ci, val in enumerate([no, name, sid, role]):
        c = row.cells[ci]
        c.text = val
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if c.paragraphs[0].runs:
            set_font(c.paragraphs[0].runs[0], size=11)

para(doc, space_after=8)
para(doc, "Advisor : Dr. Adhi Setyo Santoso, ST., MBA.", center=True, size=12, space_after=24)

para(doc, "Submitted for", center=True, size=12, space_after=2)
para(doc, "Capstone Design Project", center=True, size=12, space_after=2)
para(doc, "to Faculty of Computer Science", center=True, size=12, space_after=2)
para(doc, "President University", center=True, size=12, space_after=2)
para(doc, "2026", center=True, size=12, space_after=0)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═══════════════════════════════════════════════════════════════════
heading(doc, "TABLE OF CONTENTS", level=1, center=True, size=13)

toc_items = [
    ("STATEMENT OF ORIGINALITY", "3"),
    ("SCREENSHOT OF ZEROGPT", "4"),
    ("PART 3 – DESIGN (F300)", "5"),
    ("A. SYSTEM DESIGN", "6"),
    ("    A.1  System Architecture", "6"),
    ("    A.2  Mock-up Designs", "7"),
    ("B. HIERARCHICAL / ITERATIVE DESIGN", "9"),
    ("    B.1  Data Flow Diagram", "9"),
    ("    B.2  Interface Information Between Components", "11"),
    ("    B.3  Software Engineering Design Steps", "12"),
    ("    B.4  Component References and Libraries", "13"),
    ("    B.5  UML Diagrams", "14"),
    ("C. STANDARDS USED", "18"),
    ("D. IMPLEMENTATION AND TESTING SCENARIO", "20"),
    ("    D.1  Implementation Overview", "20"),
    ("    D.2  Testing Scenarios", "21"),
    ("REFERENCES", "27"),
]
for item, page in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(item)
    set_font(r, size=12)
    # tab stop + page number
    r2 = p.add_run(f"\t{page}")
    set_font(r2, size=12)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# STATEMENT OF ORIGINALITY
# ═══════════════════════════════════════════════════════════════════
heading(doc, "STATEMENT OF ORIGINALITY", level=1, center=True)
para(doc, space_after=6)
body(doc,
     "In my capacity as an active student at President University and as the author "
     "of the Capstone Design Project stated below:", indent=False)
para(doc, space_after=4)

# Name list
p = doc.add_paragraph()
r = p.add_run("Name")
set_font(r, size=12)
r2 = p.add_run("\t: 1. Auza Syamil Nabawi – 001202300150")
set_font(r2, size=12)

p2 = doc.add_paragraph()
r3 = p2.add_run("\t  2. Rashad Abdul Faqih – 001202300149")
set_font(r3, size=12)

p3 = doc.add_paragraph()
r4 = p3.add_run("\t  3. Atthariqul Hazam Albanna – 012202300122")
set_font(r4, size=12)

p4 = doc.add_paragraph()
r5 = p4.add_run("Faculty")
set_font(r5, size=12)
r6 = p4.add_run("\t: Computer Science")
set_font(r6, size=12)

para(doc, space_after=6)
body(doc,
     'I hereby declare that my Capstone Design Project entitled '
     '"Integrated Operational Management System for C3MR: Web Admin Portal & '
     'Telegram Mini App Ecosystem" is to the best of my knowledge and belief, an '
     'original piece of work based on sound academic principles. If there is any '
     'plagiarism detected in this final project, I am willing to be personally '
     'responsible for the consequences of these acts of plagiarism and will accept '
     'the sanctions against these acts in accordance with the rules and policies of '
     'President University.', indent=False)
para(doc, space_after=6)
body(doc,
     "I also declare that this work, either in whole or in part, has not been "
     "submitted to another university to obtain a degree.", indent=False)
para(doc, space_after=18)
para(doc, "Cikarang, 2026", center=False, size=12, space_after=18)

# Signature table
sig_tbl = doc.add_table(rows=3, cols=3)
sig_tbl.style = "Table Grid"
sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_sig = ["Signer 1", "Signer 2", "Signer 3"]
names_sig = [
    "Auza Syamil Nabawi –\n001202300150",
    "Rashad Abdul Faqih –\n001202300149",
    "Atthariqul Hazam Albanna –\n012202300122",
]
for ci, h in enumerate(headers_sig):
    c = sig_tbl.rows[0].cells[ci]
    c.text = h
    if c.paragraphs[0].runs:
        set_font(c.paragraphs[0].runs[0], size=11)

# Blank row for signatures
for ci in range(3):
    sig_tbl.rows[1].cells[ci].text = "\n\n\n"

for ci, n in enumerate(names_sig):
    c = sig_tbl.rows[2].cells[ci]
    c.text = n
    if c.paragraphs[0].runs:
        set_font(c.paragraphs[0].runs[0], size=11)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# SCREENSHOT OF ZEROGPT
# ═══════════════════════════════════════════════════════════════════
heading(doc, "SCREENSHOT OF ZEROGPT", level=1, center=True)
body(doc,
     "The entire document must be checked using ZeroGPT to verify originality. "
     "Due to the free-tier limit of 15,000 characters per submission, multiple "
     "iterations are required. Insert the ZeroGPT result screenshots below for "
     "each iteration.")
para(doc, space_after=6)
numbered(doc, "Characters 0 – 15,000")
para(doc, "[Insert ZeroGPT screenshot here]", italic=True, center=True, space_after=12)
numbered(doc, "Characters 15,001 – 30,000")
para(doc, "[Insert ZeroGPT screenshot here]", italic=True, center=True, space_after=12)
numbered(doc, "Characters 30,001 – end")
para(doc, "[Insert ZeroGPT screenshot here]", italic=True, center=True, space_after=12)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# PART 3 – DESIGN
# ═══════════════════════════════════════════════════════════════════
heading(doc, "PART 3", level=1, center=True, size=14, space_before=0)
heading(doc, "DESIGN (F300)", level=1, center=True, size=14, space_before=0)
para(doc, space_after=8)
p_consists = doc.add_paragraph()
r = p_consists.add_run("Consists of:")
set_font(r, bold=True, size=12)

for sec in ["A. SYSTEM DESIGN",
            "B. HIERARCHICAL / ITERATIVE DESIGN",
            "C. STANDARDS USED",
            "D. IMPLEMENTATION AND TESTING SCENARIO"]:
    p = doc.add_paragraph()
    r = p.add_run(f"    {sec}")
    set_font(r, bold=True, size=12)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# A. SYSTEM DESIGN
# ═══════════════════════════════════════════════════════════════════
heading(doc, "A. SYSTEM DESIGN", level=1, center=True)

h2(doc, "Overview")
body(doc,
     "C3MR (Collection Case & Customer Management Report) is a web-based field "
     "collection management system designed to streamline the workflow between "
     "field officers and managers in the debt collection process. The system "
     "follows a three-tier architecture consisting of a Presentation Layer, "
     "Business Logic Layer, and Data Storage Layer.", indent=False)

h2(doc, "System Architecture")
body(doc,
     "The system architecture is illustrated in the diagram below:", indent=False)

add_image(doc,
          os.path.join(BASE, "uml", "images", "01_system_architecture.png"),
          width=Inches(5.5),
          caption="Figure A.1  C3MR System Architecture Diagram")

para(doc, space_after=6)
body(doc, "The architecture consists of the following tiers:", indent=False)

h3(doc, "1. Client Tier (Presentation Layer)")
bullet(doc, "Web Admin Portal (React.js): A browser-based dashboard used by managers to upload "
            "collection targets (CSV), assign officers, and monitor progress in real time.")
bullet(doc, "Telegram Mini App (Web App SDK): A lightweight mobile interface embedded within "
            "Telegram, used by field officers to view assigned tasks, fill visit reports, upload "
            "photo evidence, and submit results.")
bullet(doc, "Manager Bot (Telegram Chat): A Telegram bot that allows managers to receive "
            "notifications and issue commands (e.g., reassign targets, view summaries).")

h3(doc, "2. Internet Gateway")
bullet(doc, "HTTPS / WSS Gateway (Cloudflare / Nginx): Acts as a reverse proxy and load "
            "balancer handling SSL termination, rate limiting, and request routing.")

h3(doc, "3. Application Tier (Business Logic)")
bullet(doc, "Backend API Server (Python FastAPI): Handles all API requests, validates input "
            "data via Pydantic, executes business logic (task assignment, report processing, "
            "status updates), and communicates with the database.")

h3(doc, "4. Data Tier (Storage Layer)")
bullet(doc, "Relational Database (Supabase PostgreSQL): Stores user accounts, collection "
            "targets, visit reports, officer assignments, and audit logs.")
bullet(doc, "Object Storage (Supabase Storage): Stores photo evidence from field visits.")

# ── Mock-up Designs ──
h2(doc, "Mock-up Designs")

h3(doc, "Web Admin Portal Mock-up")
add_image(doc,
          os.path.join(BASE, "uml", "images", "02_mockup_web_admin.png"),
          caption="Figure A.2  Web Admin Portal Mock-up")
para(doc, space_after=4)
body(doc, "The Web Admin Portal provides the following interface elements:", indent=False)
bullet(doc, "Navigation Tabs: Dashboard, Targets, Officers, Settings.")
bullet(doc, "Upload Targets Section: CSV file upload area with 'Choose CSV' and "
            "'Upload & Assign' buttons.")
bullet(doc, "System Summary: Four stat cards — Total Targets, Completed (green), "
            "In Progress (yellow), and Pending (red).")
bullet(doc, "Target List Table: Sortable/filterable table showing Target ID, Customer name, "
            "Address, Amount Due, assigned Officer, Status (Completed / In Progress / Pending), "
            "and an Actions column (View or Assign).")

h3(doc, "Telegram Mini App Mock-up")
add_image(doc,
          os.path.join(BASE, "uml", "images", "03_mockup_telegram_miniapp.png"),
          caption="Figure A.3  Telegram Mini App Mock-up")
para(doc, space_after=4)
body(doc, "The Telegram Mini App provides the following interface elements:", indent=False)
bullet(doc, 'Header: App name "C3MR Field App" with back navigation.')
bullet(doc, "Task Detail: Displays the current task with customer name, address, phone "
            "number, and amount due (highlighted in red).")
bullet(doc, "Payment Status Dropdown: Selectable options — Promise to Pay, Paid, Refused, "
            "Not Home, Partial Payment.")
bullet(doc, "Notes Field: Scrollable text area for the officer to enter visit notes.")
bullet(doc, "Upload Photo Evidence Button: Triggers the device camera or gallery.")
bullet(doc, "Submit Report Button: Sends the completed report to the backend API.")

h3(doc, "System Dashboard Screenshot")
add_image(doc,
          os.path.join(BASE, "Gemini_Generated_Image_axio0iaxio0iaxio.png"),
          width=Inches(5.5),
          caption="Figure A.4  C3MR Executive Overview Dashboard (Rendered Design)")
para(doc, space_after=6)
add_image(doc,
          os.path.join(BASE, "Gemini_Generated_Image_l1qpc6l1qpc6l1qp (1).png"),
          width=Inches(5.5),
          caption="Figure A.5  C3MR User Management Page (Rendered Design)")

page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# B. HIERARCHICAL / ITERATIVE DESIGN
# ═══════════════════════════════════════════════════════════════════
heading(doc, "B. HIERARCHICAL / ITERATIVE DESIGN", level=1, center=True)

h2(doc, "1. Data Flow Diagram")

h3(doc, "Level 0 – Context Diagram")
body(doc,
     "The Level 0 (Context) Diagram shows the top-level data flows between the "
     "three external entities and the C3MR system:", indent=False)

dfd_text = (
    "+-----------------+         +-------------------+         +-------------------+\n"
    "|  Field Officer  | ------> |                   | <------ |     Manager       |\n"
    "| (Telegram Mini  |  Submit |      C3MR         |  Upload | (Web Admin Portal |\n"
    "|     App)        |  Report |      System       |  CSV /  |  / Manager Bot)   |\n"
    "|                 | <------ |                   | ------> |                   |\n"
    "|                 |  Task   |                   |  Reports|                   |\n"
    "|                 |  List   |                   |  & Stats|                   |\n"
    "+-----------------+         +-------------------+         +-------------------+\n"
    "                                    |     ^\n"
    "                                    v     |\n"
    "                            +-------------------+\n"
    "                            |  Supabase         |\n"
    "                            |  (PostgreSQL +    |\n"
    "                            |   Storage)        |\n"
    "                            +-------------------+"
)
p_dfd = doc.add_paragraph()
r_dfd = p_dfd.add_run(dfd_text)
set_font(r_dfd, name="Courier New", size=9)

h3(doc, "Level 1 – Major Processes")
for proc, desc in [
    ("1.0 Authentication & Authorization",
     "Handles user login (Supabase Auth for admin, Telegram user validation for officers), "
     "session management, and role-based access control."),
    ("2.0 Target Management",
     "Manages CSV upload, parsing, validation, and assignment of collection targets to field officers."),
    ("3.0 Report Management",
     "Processes field visit reports including form data validation, photo upload to object storage, "
     "and database insertion."),
    ("4.0 Officer Management",
     "CRUD operations for field officer records, including linking Telegram user IDs."),
    ("5.0 Dashboard & Analytics",
     "Aggregates data for the admin dashboard: total targets, completion rates, officer performance."),
    ("6.0 Notification Service",
     "Sends task assignment notifications and status updates via the Telegram Bot API."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{proc}: ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(desc)
    set_font(r2, size=12)

h3(doc, "Level 2 – Report Submission Process (Detailed)")
seq_text = (
    "Field Officer          Mini App (UI)         FastAPI Backend       Supabase DB\n"
    "     |                      |                        |                    |\n"
    "     |-- Fill form & ------>|                        |                    |\n"
    "     |   photo              |-- POST /api/reports -->|                    |\n"
    "     |                      |   (JSON + Image)       |-- Upload image --->|\n"
    "     |                      |                        |<-- Image URL ------|\n"
    "     |                      |                        |-- INSERT report -->|\n"
    "     |                      |                        |   & UPDATE target  |\n"
    "     |                      |                        |<-- Success --------|\n"
    "     |                      |<-- HTTP 200 OK --------|                    |\n"
    "     |<-- Show 'Success' ---|                        |                    |"
)
p_seq = doc.add_paragraph()
r_seq = p_seq.add_run(seq_text)
set_font(r_seq, name="Courier New", size=9)

h2(doc, "2. Interface Information Between Components")

interfaces = [
    ("Client ↔ Gateway",
     "Protocol: HTTPS (TLS 1.3); Data format: JSON; Authentication: Bearer token."),
    ("Gateway ↔ Backend API",
     "Protocol: HTTP (internal); Headers forwarded: Authorization, X-Forwarded-For."),
    ("Backend API ↔ Database",
     "Protocol: PostgreSQL wire (TCP); Interface: supabase-py SDK; Parameterized SQL queries."),
    ("Backend API ↔ Object Storage",
     "Protocol: HTTPS; Interface: Supabase Storage API; Multipart file upload; Returns public URL."),
    ("Backend API ↔ Telegram Bot API",
     "Protocol: HTTPS; HTTP POST with JSON; sendMessage / sendPhoto endpoints."),
]
for title_i, desc_i in interfaces:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{title_i}: ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(desc_i)
    set_font(r2, size=12)

h2(doc, "3. Software Engineering Design Steps")
body(doc,
     "The project follows an Agile (Iterative) development methodology:", indent=False)
steps = [
    ("Requirements Gathering",
     "Identify stakeholder needs, define user stories, and prioritize features in a product backlog."),
    ("System Design",
     "Create architecture diagrams, define API contracts (OpenAPI spec), design database schema, "
     "and produce UI mock-ups."),
    ("Sprint-based Implementation",
     "Develop features incrementally in 2-week sprints, starting with authentication, target upload, "
     "and report submission."),
    ("Testing",
     "Unit tests (pytest), integration tests (API endpoint testing), and UAT with stakeholders."),
    ("Deployment",
     "Deploy backend to cloud server, configure Supabase project, register Telegram bot, "
     "and publish Mini App."),
    ("Review & Iteration",
     "Collect feedback, fix bugs, and iterate on features."),
]
for i, (st, sd) in enumerate(steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{i}. {st}: ")
    set_font(r1, bold=True, size=12)
    r2 = p.add_run(sd)
    set_font(r2, size=12)

h2(doc, "4. Component References and Libraries")

h3(doc, "Frontend – Web Admin Portal")
for lib in ["React.js (v18.x) – UI framework",
            "React Router (v6.x) – Client-side routing",
            "Axios – HTTP client for API calls",
            "Tailwind CSS – Utility-first CSS framework",
            "Papa Parse – CSV parsing library",
            "Chart.js / Recharts – Dashboard chart rendering"]:
    bullet(doc, lib)

h3(doc, "Frontend – Telegram Mini App")
for lib in ["Telegram Web App SDK – Native Telegram integration",
            "HTML5 / CSS3 / JavaScript – Core web technologies",
            "Fetch API – HTTP client for API calls"]:
    bullet(doc, lib)

h3(doc, "Backend")
for lib in ["Python (v3.11+) – Programming language",
            "FastAPI (v0.100+) – Web framework",
            "Pydantic (v2.x) – Data validation and serialization",
            "supabase-py – Supabase Python client SDK",
            "python-telegram-bot – Telegram Bot API wrapper",
            "Uvicorn – ASGI server",
            "python-multipart – File upload handling"]:
    bullet(doc, lib)

h3(doc, "Database & Storage")
for lib in ["Supabase – Backend-as-a-Service platform",
            "PostgreSQL (v15) – Relational database",
            "Supabase Storage – S3-compatible object storage",
            "Supabase Auth – Authentication service"]:
    bullet(doc, lib)

h2(doc, "5. UML Diagrams")

for img_path, caption in [
    (os.path.join(BASE, "uml", "images", "05_sequence_diagram.png"),
     "Figure B.1  Sequence Diagram – Report Submission Flow"),
    (os.path.join(BASE, "uml", "images", "04_component_diagram.png"),
     "Figure B.2  Component Diagram"),
]:
    add_image(doc, img_path, caption=caption)
    para(doc, space_after=8)

# ERD text
h3(doc, "Entity Relationship Diagram")
erd_text = (
    "+-------------+       +-------------+       +-------------+\n"
    "|   users     |       |   targets   |       |   reports   |\n"
    "+-------------+       +-------------+       +-------------+\n"
    "| PK id       |<--+   | PK id       |<---+  | PK id       |\n"
    "| telegram_id |   |   | customer_   |    |  | FK target_id|\n"
    "| name        |   |   |   name      |    |  | FK officer_ |\n"
    "| role        |   +---| FK assigned_|    +--|   id        |\n"
    "| created_at  |       |   officer   |       | payment_    |\n"
    "+-------------+       | address     |       |   status    |\n"
    "                      | phone       |       | notes       |\n"
    "                      | amount_due  |       | photo_url   |\n"
    "                      | status      |       | submitted_at|\n"
    "                      | created_at  |       +-------------+\n"
    "                      +-------------+"
)
p_erd = doc.add_paragraph()
r_erd = p_erd.add_run(erd_text)
set_font(r_erd, name="Courier New", size=9)

body(doc, "Relationships:", indent=False)
for rel in [
    "A User (officer) can be assigned to many Targets (1:N).",
    "A Target can have many Reports (1:N) – e.g., multiple visit attempts.",
    "A Report belongs to one Target and one User (officer).",
    "An UploadBatch is created by one User (manager) and generates many Targets.",
]:
    bullet(doc, rel)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# C. STANDARDS USED
# ═══════════════════════════════════════════════════════════════════
heading(doc, "C. STANDARDS USED", level=1, center=True)

standards = [
    ("Data Format",       "JSON (RFC 8259)",           "All API request/response bodies use JSON."),
    ("API Architecture",  "REST (RESTful API)",         "Resource-based URLs, HTTP methods, status codes."),
    ("API Documentation", "OpenAPI 3.0",                "FastAPI auto-generates OpenAPI 3.0 spec (Swagger UI)."),
    ("Data Validation",   "Pydantic v2",                "Strict type checking and custom validators on all inputs."),
    ("Authentication",    "JWT (RFC 7519)",             "Supabase Auth issues JWTs upon login."),
    ("Encryption",        "TLS 1.3 (RFC 8446)",         "All client-server communication encrypted via HTTPS."),
    ("Database",          "SQL (PostgreSQL 15)",        "Standard SQL with PostgreSQL extensions."),
    ("Modeling",          "UML 2.5",                   "Use case, class, sequence, activity, component, ERD."),
    ("Diagram Tool",      "PlantUML",                   "All UML diagrams authored in PlantUML markup → PNG."),
    ("Code Formatter",    "Black (Python)",             "Line length 88 (default settings)."),
    ("Code Formatter",    "Prettier (JS/TS)",           "Frontend JavaScript/TypeScript formatting."),
    ("Linter",            "Ruff (Python)",              "Fast, comprehensive style and error checking."),
    ("Linter",            "ESLint (JS/TS)",             "Frontend code linting with recommended rules."),
    ("CSS Framework",     "Tailwind CSS v3",            "Utility-first CSS for consistent, responsive styling."),
    ("Version Control",   "Git",                        "GitHub as the remote repository host."),
    ("File Upload",       "Multipart/form-data (RFC 7578)", "Photo evidence uploads."),
    ("CSV Format",        "RFC 4180",                   "Target data CSV files."),
    ("Bot API",           "Telegram Bot API v7",        "Official Bot API for messages and Mini App integration."),
    ("Character Encoding","UTF-8",                      "All text data throughout the system."),
    ("Date/Time Format",  "ISO 8601",                   "All timestamps (e.g., 2024-01-15T10:30:00Z)."),
]

simple_table(
    doc,
    headers=["Category", "Standard / Technology", "Description"],
    rows=standards,
    col_widths=[Cm(3.5), Cm(5), Cm(8)],
)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# D. IMPLEMENTATION AND TESTING SCENARIO
# ═══════════════════════════════════════════════════════════════════
heading(doc, "D. IMPLEMENTATION AND TESTING SCENARIO", level=1, center=True)

h2(doc, "Implementation Overview")
body(doc, "The system is implemented in an iterative manner across four phases:",
     indent=False)

phases = [
    ("Phase 1 – Backend Foundation",
     ["Set up FastAPI project with Uvicorn as the ASGI server.",
      "Configure Supabase project (PostgreSQL database, Storage buckets, Auth).",
      "Define database schema and create tables (users, targets, reports, upload_batches).",
      "Implement authentication middleware using Supabase JWT verification.",
      "Implement core API endpoints: user registration, target CRUD, report submission."]),
    ("Phase 2 – Web Admin Portal",
     ["Set up React.js project with Tailwind CSS.",
      "Implement login page with Supabase Auth integration.",
      "Build dashboard page with summary statistics and charts.",
      "Build target management page with CSV upload and officer assignment.",
      "Build report viewing page with filtering and search."]),
    ("Phase 3 – Telegram Mini App & Bot",
     ["Register Telegram bot via BotFather and configure webhook.",
      "Implement Mini App using Telegram Web App SDK.",
      "Build task list view and report submission form.",
      "Implement photo capture and upload functionality.",
      "Implement Manager Bot commands for notifications and quick actions."]),
    ("Phase 4 – Integration & Testing",
     ["End-to-end integration testing across all components.",
      "User Acceptance Testing (UAT) with stakeholders.",
      "Performance testing and optimization.",
      "Bug fixes and final adjustments."]),
]
for phase_name, phase_items in phases:
    p = doc.add_paragraph()
    r = p.add_run(phase_name)
    set_font(r, bold=True, size=12)
    for item in phase_items:
        bullet(doc, item)

h2(doc, "Testing Scenarios")

# ── Scenario tables ──
scenarios = [
    ("Scenario 1: Manager Login (Web Admin Portal)",
     ["Step", "Action", "Expected Result", "Alternative"],
     [
         ("1", "Manager navigates to the Web Admin Portal URL.",
          "Login page is displayed with email and password fields.",
          "If server is unreachable, show 'Unable to connect' error."),
         ("2", "Manager enters valid email and password, clicks 'Login'.",
          "System authenticates via Supabase Auth, redirects to Dashboard. JWT stored in local storage.",
          "If credentials invalid, show 'Invalid email or password'. User stays on login page."),
         ("3", "Manager enters email but leaves password empty.",
          "Client-side validation shows 'Password is required'. No API call is made.", "–"),
         ("4", "Manager's JWT token expires during session.",
          "System detects 401 response, redirects to login with 'Session expired' message.", "–"),
     ]),
    ("Scenario 2: Upload CSV Targets (Web Admin Portal)",
     ["Step", "Action", "Expected Result", "Alternative"],
     [
         ("1", "Manager clicks 'Choose CSV' button.",
          "File picker opens, filtered to .csv files.", "–"),
         ("2", "Manager selects valid CSV with correct columns.",
          "File name displayed. Preview of first 5 rows shown.",
          "If CSV has wrong columns, show format error."),
         ("3", "Manager clicks 'Upload & Assign'.",
          "System parses CSV, creates target records. Success: '150 targets uploaded.'",
          "If any row has invalid data, show row-level error. Upload rejected."),
         ("4", "Manager uploads CSV with 0 data rows.",
          "Show 'CSV file contains no data rows' error.", "–"),
         ("5", "Manager uploads a non-CSV file (e.g., .xlsx).",
          "Client-side validation: 'Only .csv files are accepted.'", "–"),
     ]),
    ("Scenario 3: Assign Targets to Officers (Web Admin Portal)",
     ["Step", "Action", "Expected Result", "Alternative"],
     [
         ("1", "Manager selects unassigned targets.",
          "Selected targets highlighted. 'Assign' button becomes active.", "–"),
         ("2", "Manager clicks 'Assign' and selects officer.",
          "Targets assigned. Status → 'In Progress'. Telegram notification sent to officer.",
          "If officer is at max capacity, show warning."),
         ("3", "Manager attempts to reassign an already-assigned target.",
          "Confirmation dialog appears. On confirm, reassignment proceeds.",
          "If manager cancels, no changes made."),
     ]),
    ("Scenario 4: Field Officer Views Task List (Telegram Mini App)",
     ["Step", "Action", "Expected Result", "Alternative"],
     [
         ("1", "Field Officer opens the Mini App from Telegram.",
          "System validates Telegram user ID. Task list displayed.",
          "If ID not registered, show 'Contact your manager' message."),
         ("2", "Task list loads with assigned targets.",
          "Each task shows customer name, address snippet, amount due, and status badge. "
          "Sorted by status (Pending first).",
          "If no tasks, show 'No tasks assigned. Check back later.'"),
         ("3", "Officer taps on a task.",
          "Task detail view opens with full customer info (name, address, phone, amount due).", "–"),
     ]),
    ("Scenario 5: Field Officer Submits Report (Telegram Mini App)",
     ["Step", "Action", "Expected Result", "Alternative"],
     [
         ("1", "Officer opens task detail and taps 'Submit Report'.",
          "Report form displayed with Payment Status dropdown, Notes, and Upload button.", "–"),
         ("2", "Officer selects payment status (Promise to Pay / Paid / Refused / Not Home / Partial Payment).",
          "Selection recorded. Dropdown shows selected value.", "–"),
         ("3", "Officer types notes.",
          "Text recorded in notes field.", "–"),
         ("4", "Officer taps 'Upload Photo Evidence'.",
          "Camera or gallery opens.",
          "If permission denied, show 'Please grant camera access.'"),
         ("5", "Officer captures/selects a photo.",
          "Photo thumbnail preview displayed.",
          "If photo > 10 MB, show size limit error."),
         ("6", "Officer taps 'Submit Report' with all fields filled.",
          "POST /api/reports sent. Photo uploaded. Report inserted. Target → 'Completed'. "
          "Success message shown. Officer returned to task list.",
          "If network error, show retry message. Form data preserved."),
         ("7", "Officer taps 'Submit Report' without payment status.",
          "Validation error: 'Please select a payment status.' Form not submitted.", "–"),
         ("8", "Officer taps 'Submit Report' without photo.",
          "Validation error: 'Please upload photo evidence.' Form not submitted.", "–"),
     ]),
    ("Scenario 6: Manager Views Dashboard (Web Admin Portal)",
     ["Step", "Action", "Expected Result", "Alternative"],
     [
         ("1", "Manager navigates to Dashboard page.",
          "Dashboard loads with 4 summary cards (Total Targets, Completed, In Progress, Pending) "
          "and completion rate chart.",
          "If no data, show 'Upload targets to get started.'"),
         ("2", "Manager views Target List table.",
          "Table shows: ID, Customer, Address, Amount Due, Officer, Status, Actions. "
          "Supports sorting and filtering by status.", "–"),
         ("3", "Manager clicks on a completed target row.",
          "Target detail view shows customer info and submitted report "
          "(payment status, notes, photo evidence).", "–"),
     ]),
    ("Scenario 7: Manager Bot Notifications (Telegram)",
     ["Step", "Action", "Expected Result", "Alternative"],
     [
         ("1", "A field officer submits a report.",
          "Manager receives Telegram notification: "
          "'📋 New Report: [Customer] – [Payment Status] by [Officer].'",
          "If bot blocked by manager, notification fails silently. Log error."),
         ("2", "Manager sends /summary to the bot.",
          "Bot responds with total targets, completed, pending count, and completion %.",
          "If no targets, bot replies 'No targets found.'"),
     ]),
]

for sc_title, sc_headers, sc_rows in scenarios:
    h3(doc, sc_title)
    simple_table(doc, sc_headers, sc_rows,
                 col_widths=[Cm(1.2), Cm(4.5), Cm(5.8), Cm(4.5)])
    para(doc, space_after=8)

page_break(doc)

# ═══════════════════════════════════════════════════════════════════
# REFERENCES
# ═══════════════════════════════════════════════════════════════════
heading(doc, "REFERENCES", level=1, center=True)

refs = [
    "FastAPI Documentation. (n.d.). Retrieved from https://fastapi.tiangolo.com/",
    "Supabase Documentation. (n.d.). Retrieved from https://supabase.com/docs",
    "Telegram Bot API Documentation. (n.d.). Retrieved from https://core.telegram.org/bots/api",
    "Telegram Mini Apps Documentation. (n.d.). Retrieved from https://core.telegram.org/bots/webapps",
    "React.js Documentation. (n.d.). Retrieved from https://react.dev/",
    "PlantUML Documentation. (n.d.). Retrieved from https://plantuml.com/",
    "Pydantic Documentation. (n.d.). Retrieved from https://docs.pydantic.dev/",
    "RFC 8259 – The JavaScript Object Notation (JSON) Data Interchange Format. (2017). "
    "Retrieved from https://datatracker.ietf.org/doc/html/rfc8259",
    "RFC 7519 – JSON Web Token (JWT). (2015). "
    "Retrieved from https://datatracker.ietf.org/doc/html/rfc7519",
    "Tailwind CSS Documentation. (n.d.). Retrieved from https://tailwindcss.com/docs",
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent     = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after     = Pt(6)
    r = p.add_run(ref)
    set_font(r, size=12)

# ── Save ─────────────────────────────────────────────────────────────────────
out = os.path.join(BASE, "Form3_Design.docx")
doc.save(out)
print(f"Saved: {out}")
